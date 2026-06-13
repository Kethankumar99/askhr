import os
import uuid
from sqlalchemy.orm import Session
from app.models.document import Document
import chromadb

# Initialize ChromaDB
CHROMA_PATH = "./data/chroma_db"
os.makedirs(CHROMA_PATH, exist_ok=True)

chroma_client = chromadb.PersistentClient(path=CHROMA_PATH)

def get_or_create_collection(company_email: str):
    """Get or create ChromaDB collection for a company"""
    # Clean name for collection
    safe_name = company_email.replace('@', '_').replace('.', '_').replace('-', '_')
    collection_name = f"company_{safe_name}"
    
    try:
        collection = chroma_client.get_collection(collection_name)
    except:
        collection = chroma_client.create_collection(collection_name)
    
    return collection

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """Extract text from different file types"""
    import io
    
    text = ""
    filename_lower = filename.lower()
    
    try:
        if filename_lower.endswith('.txt'):
            text = file_content.decode('utf-8', errors='ignore')
        
        elif filename_lower.endswith('.pdf'):
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        elif filename_lower.endswith('.docx'):
            import docx
            doc = docx.Document(io.BytesIO(file_content))
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
        
        elif filename_lower.endswith('.csv'):
            text = file_content.decode('utf-8', errors='ignore')
        
        else:
            raise ValueError(f"Unsupported file type: {filename}")
    
    except Exception as e:
        raise Exception(f"Error extracting text: {str(e)}")
    
    return text.strip()

def chunk_text_simple(text: str, chunk_size: int = 500, chunk_overlap: int = 50) -> list:
    """Split text into overlapping chunks"""
    
    # Split by sentences first
    sentences = []
    current = ""
    
    for char in text:
        current += char
        if char in ['.', '!', '?', '\n'] and len(current) > 20:
            sentences.append(current.strip())
            current = ""
    
    if current.strip():
        sentences.append(current.strip())
    
    # Create chunks
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        if len(current_chunk) + len(sentence) <= chunk_size:
            current_chunk += sentence + " "
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = sentence + " "
    
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    # If no chunks created, use simple split
    if not chunks:
        words = text.split()
        for i in range(0, len(words), chunk_size):
            chunk = ' '.join(words[i:i+chunk_size])
            if chunk:
                chunks.append(chunk)
    
    return chunks

def upload_document(db: Session, file_content: bytes, filename: str, company_email: str) -> dict:
    """Upload document to DB + Vector Store"""
    
    try:
        print(f"\n📄 Processing: {filename}")
        
        # Extract text
        text = extract_text_from_file(file_content, filename)
        
        if not text:
            return {"success": False, "message": "No text found in document"}
        
        print(f"📝 Extracted {len(text)} characters")
        
        # Chunk text
        chunks = chunk_text_simple(text)
        print(f"✂️ Created {len(chunks)} chunks")
        
        # Save to SQLite
        unique_filename = f"{uuid.uuid4()}_{filename}"
        
        new_doc = Document(
            filename=unique_filename,
            original_name=filename,
            file_type=filename.split('.')[-1].lower(),
            company_email=company_email,
            chunk_count=len(chunks)
        )
        
        db.add(new_doc)
        db.commit()
        db.refresh(new_doc)
        
        # Save to ChromaDB
        collection = get_or_create_collection(company_email)
        
        chunk_ids = [f"doc_{new_doc.id}_chunk_{i}" for i in range(len(chunks))]
        metadatas = [{"document_id": new_doc.id, "filename": filename} for _ in chunks]
        
        collection.add(
            documents=chunks,
            ids=chunk_ids,
            metadatas=metadatas
        )
        
        print(f"✅ Document saved! ID: {new_doc.id}")
        
        return {
            "success": True,
            "message": f"Document uploaded. {len(chunks)} chunks created.",
            "document": {
                "id": new_doc.id,
                "filename": filename,
                "chunk_count": len(chunks),
                "file_type": new_doc.file_type,
                "created_at": str(new_doc.created_at)
            }
        }
    
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        return {"success": False, "message": str(e)}

def get_all_documents(db: Session, company_email: str) -> dict:
    """Get all documents for company"""
    
    documents = db.query(Document).filter(
        Document.company_email == company_email
    ).order_by(Document.created_at.desc()).all()
    
    doc_list = []
    for doc in documents:
        doc_list.append({
            "id": doc.id,
            "filename": doc.original_name,
            "file_type": doc.file_type,
            "chunk_count": doc.chunk_count,
            "created_at": str(doc.created_at)
        })
    
    return {"success": True, "count": len(doc_list), "documents": doc_list}

def delete_document(db: Session, document_id: int, company_email: str) -> dict:
    """Delete document from DB + Vector Store"""
    
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.company_email == company_email
    ).first()
    
    if not doc:
        return {"success": False, "message": "Document not found"}
    
    try:
        # Delete from Vector DB
        collection = get_or_create_collection(company_email)
        chunk_ids = [f"doc_{doc.id}_chunk_{i}" for i in range(doc.chunk_count)]
        
        try:
            collection.delete(ids=chunk_ids)
        except:
            pass
        
        # Delete from SQLite
        db.delete(doc)
        db.commit()
        
        return {"success": True, "message": f"Document '{doc.original_name}' deleted"}
    
    except Exception as e:
        return {"success": False, "message": str(e)}